"""
小红书工具箱 — 统一 Web 入口
包含：博主分析、帖子深挖、写作库
一键启动：python app.py
"""

import os
import sys
import re
import json
import time
import queue
import threading
import subprocess
from datetime import datetime
from pathlib import Path

os.chdir(os.path.dirname(os.path.abspath(__file__)))
# 嵌入版 Python 不会自动把脚本目录加入 sys.path，需要手动添加
if os.getcwd() not in sys.path:
    sys.path.insert(0, os.getcwd())

from flask import Flask, render_template, request, jsonify, Response, send_file, send_from_directory

app = Flask(__name__)
app.config["TEMPLATES_AUTO_RELOAD"] = True

# ── 全局 CORS 头（兼容浏览器限制） ──
@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    return response

# ── 路径常量 ──
ROOT = Path(".").resolve()
TOOLS_DIR = ROOT / "工具类"
MATERIAL_DIR = ROOT / "素材库"
OUTPUT_DIR = ROOT / "output"
SPIDER_DIR = TOOLS_DIR / "博主蒸馏" / "spider_xhs"
DATA_DIR = TOOLS_DIR / "博主蒸馏" / "data"

# 素材库子目录
MATERIAL_BLOGGER_DIR = MATERIAL_DIR / "博主风格"  # 已分析的博主风格（deep_analyze产出）
MATERIAL_SELF_DIR = MATERIAL_DIR / "自我参数"    # 自己的写作参数

DATA_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
MATERIAL_SELF_DIR.mkdir(parents=True, exist_ok=True)
MATERIAL_BLOGGER_DIR.mkdir(parents=True, exist_ok=True)

# ── 全局任务管理 ──
_tasks: dict[str, dict] = {}
_task_lock = threading.Lock()
_running_procs: dict[str, subprocess.Popen] = {}
_procs_lock = threading.Lock()
_task_stop_flags: dict[str, threading.Event] = {}
_stop_flags_lock = threading.Lock()


def _next_id() -> str:
    return f"task_{int(time.time())}_{len(_tasks)}"


def _get_stop_event(task_id: str) -> threading.Event:
    with _stop_flags_lock:
        if task_id not in _task_stop_flags:
            _task_stop_flags[task_id] = threading.Event()
        return _task_stop_flags[task_id]


def _cleanup_stop_event(task_id: str):
    with _stop_flags_lock:
        _task_stop_flags.pop(task_id, None)


def _is_stopped(task_id: str) -> bool:
    with _stop_flags_lock:
        ev = _task_stop_flags.get(task_id)
        return ev is not None and ev.is_set()


def _run_subprocess(cmd: list, log_queue: queue.Queue, task_id: str | None = None, cwd=None, on_done=None):
    """在子线程中运行命令，输出逐行写入 log_queue。task_id 不为 None 时自动更新任务状态。on_done(task_id, exit_code) 在任务完成后调用"""
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    env["PYTHONIOENCODING"] = "utf-8"
    exit_code = 1
    try:
        if task_id is not None and _is_stopped(task_id):
            log_queue.put("[STOP] 任务已被用户终止，跳过启动\n")
            return
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            cwd=cwd,
            env=env,
        )
        if task_id is not None:
            with _procs_lock:
                _running_procs[task_id] = proc
        for line in iter(proc.stdout.readline, b""):
            text = line.decode("utf-8", errors="replace")
            log_queue.put(text)
        proc.wait()
        exit_code = proc.returncode
        log_queue.put(f"__EXIT_CODE__:{exit_code}\n")
    except Exception as e:
        log_queue.put(f"__EXIT_CODE__:1\n__ERROR__:{e}\n")
    finally:
        if task_id is not None:
            _cleanup_stop_event(task_id)
            with _procs_lock:
                _running_procs.pop(task_id, None)
            with _task_lock:
                if task_id in _tasks:
                    _tasks[task_id]["done"] = True
                    _tasks[task_id]["exit_code"] = exit_code
        if on_done is not None:
            try:
                on_done(task_id, exit_code)
            except Exception as e:
                print(f"[callback] on_done error: {e}")
        log_queue.put(None)


# =============================================================
#  工具函数
# =============================================================

def _find_result_json():
    files = sorted(DATA_DIR.glob("*_notes_details.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    return str(files[0]) if files else None

def _find_analysis_json():
    files = sorted(DATA_DIR.glob("*_analysis.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    return str(files[0]) if files else None

def _find_deep_output(prefix=""):
    search_roots = [OUTPUT_DIR]
    if MATERIAL_SELF_DIR.is_dir():
        search_roots.append(MATERIAL_SELF_DIR)
    if MATERIAL_BLOGGER_DIR.is_dir():
        search_roots.append(MATERIAL_BLOGGER_DIR)
    dm, tm, bk, fm, tp = [], [], [], [], []
    for root in search_roots:
        dm.extend(root.rglob(f"{prefix}*_数据底稿.md"))
        tm.extend(root.rglob(f"{prefix}*_AI蒸馏任务.md"))
        bk.extend(root.rglob(f"{prefix}*_博主深度拆解.md"))
        fm.extend(root.rglob(f"{prefix}*_内容公式总结.md"))
        tp.extend(root.rglob(f"{prefix}*_选题素材库.md"))
    # 按修改时间排序，取最新的
    for lst in [dm, tm, bk, fm, tp]:
        lst.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    return {
        "data_draft": str(dm[0]) if dm else None,
        "ai_task": str(tm[0]) if tm else None,
        "deep_breakdown": str(bk[0]) if bk else None,
        "formula_summary": str(fm[0]) if fm else None,
        "topic_library": str(tp[0]) if tp else None,
    }

def _tree_walk(dir_path: Path, prefix=""):
    """递归列出目录树"""
    items = []
    try:
        entries = sorted(dir_path.iterdir(), key=lambda e: (not e.is_dir(), e.name))
    except PermissionError:
        return items
    for entry in entries:
        name = entry.name
        if name.startswith("."):
            continue
        if entry.is_dir():
            children = _tree_walk(entry, f"{prefix}{name}/")
            items.append({"name": name, "type": "dir", "path": str(entry), "children": children})
        else:
            size = entry.stat().st_size
            items.append({"name": name, "type": "file", "path": str(entry), "size": size})
    return items


# =============================================================
#  博主蒸馏 — 路由
# =============================================================

def _run_pipeline(params: dict, log_queue: queue.Queue, task_id: str):
    """一键管道：采集 → 分析 → 蒸馏，依次执行，通过 SSE 推送阶段进度"""
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    env["PYTHONIOENCODING"] = "utf-8"
    exit_code = 0
    result_files = {}

    def _run_cmd(cmd, cwd=None):
        nonlocal exit_code
        if _is_stopped(task_id):
            log_queue.put("[STOP] 任务已被用户终止，跳过此阶段\n")
            return -1
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, cwd=cwd, env=env)
        with _procs_lock:
            _running_procs[task_id] = proc
        for line in iter(proc.stdout.readline, b""):
            text = line.decode("utf-8", errors="replace")
            log_queue.put(text)
        proc.wait()
        with _procs_lock:
            _running_procs.pop(task_id, None)
        return proc.returncode

    try:
        # ── Phase 1: 采集 ──
        log_queue.put("__PHASE__:1\n")
        python = sys.executable
        script = str(TOOLS_DIR / "博主蒸馏" / "scripts" / "spider_xhs_adapter.py")
        cmd = [python, script, "--output", str(DATA_DIR), "--max-notes", str(params["max_notes"])]
        if params.get("url"):
            cmd.extend(["--url", params["url"]])
        elif params.get("mode") == "self":
            cmd.append("--self")
        elif params.get("user_id"):
            cmd.extend(["--user-id", params["user_id"]])
        else:
            log_queue.put("[ERROR] 缺少博主 URL 或用户 ID\n")
            exit_code = 1
            return

        rc = _run_cmd(cmd)
        if rc != 0:
            log_queue.put(f"[ERROR] 采集阶段失败 (exit={rc})\n")
            exit_code = rc
            return

        details_file = _find_result_json()
        if details_file:
            result_files["details_json"] = details_file
            try:
                with open(details_file, "r", encoding="utf-8") as f:
                    raw = json.load(f)
                if isinstance(raw, dict) and raw.get("blogger", {}).get("nickname"):
                    params["blogger_name"] = raw["blogger"]["nickname"]
                raw_details = raw.get("details", []) if isinstance(raw, dict) else raw
                bad_items = [i for i, item in enumerate(raw_details) if not isinstance(item, dict)]
                if bad_items:
                    log_queue.put(f"[WARN] 数据中有 {len(bad_items)} 个非字典项（索引: {bad_items[:5]}），将跳过\n")
            except Exception as e:
                log_queue.put(f"[WARN] 数据文件读取异常: {e}\n")
        log_queue.put("__PHASE_DONE__:1\n")

        # ── Phase 2: 分析 ──
        log_queue.put("__PHASE__:2\n")
        if not details_file:
            log_queue.put("[ERROR] 未找到采集数据文件，跳过分析\n")
            exit_code = 1
            return

        cmd = [python, str(TOOLS_DIR / "博主蒸馏" / "scripts" / "analyze.py"), details_file, "-o", str(DATA_DIR)]
        rc = _run_cmd(cmd)
        if rc != 0:
            log_queue.put(f"[ERROR] 分析阶段失败 (exit={rc})\n")
            exit_code = rc
            return

        analysis_file = _find_analysis_json()
        if analysis_file:
            result_files["analysis_json"] = analysis_file
        log_queue.put("__PHASE_DONE__:2\n")

        # ── Phase 3: 蒸馏 ──
        log_queue.put("__PHASE__:3\n")
        blogger_name = params.get("blogger_name", "")
        if not blogger_name and analysis_file:
            blogger_name = Path(analysis_file).name.replace("_analysis.json", "")
        if not blogger_name:
            blogger_name = "unknown"

        if not analysis_file:
            log_queue.put("[ERROR] 未找到分析文件，跳过蒸馏\n")
            exit_code = 1
            return

        details_files = sorted(DATA_DIR.glob("*_notes_details.json"), key=lambda f: f.stat().st_mtime, reverse=True)
        details_for_deep = str(details_files[0]) if details_files else ""

        mode = params.get("deep_mode", "A")
        # 根据模式输出到对应的素材目录
        deep_out_dir = str(MATERIAL_SELF_DIR) if mode == "B" else str(MATERIAL_BLOGGER_DIR)
        os.makedirs(deep_out_dir, exist_ok=True)
        cmd = [
            python, str(TOOLS_DIR / "博主蒸馏" / "scripts" / "deep_analyze.py"),
            analysis_file, blogger_name,
            "-o", deep_out_dir,
            "--details", details_for_deep,
            "--mode", mode,
        ]
        rc = _run_cmd(cmd)
        if rc != 0:
            log_queue.put(f"[ERROR] 蒸馏阶段失败 (exit={rc})\n")
            exit_code = rc
            return

        deep_out = _find_deep_output()
        result_files.update(deep_out)
        log_queue.put("__PHASE_DONE__:3\n")

    except Exception as e:
        log_queue.put(f"[ERROR] 管道异常: {e}\n")
        exit_code = 1
    finally:
        _cleanup_stop_event(task_id)
        with _task_lock:
            if task_id in _tasks:
                _tasks[task_id]["done"] = True
                _tasks[task_id]["exit_code"] = exit_code
                _tasks[task_id]["result_files"] = result_files
        log_queue.put(None)


@app.route("/api/one-click-analyze", methods=["POST", "GET"])
def one_click_analyze():
    """一键分析：采集 → 分析 → 蒸馏，自动串联"""
    if request.method == "GET":
        data = request.args
    else:
        data = request.get_json() or request.form

    url = (data.get("url") or "").strip()
    user_id = (data.get("user_id") or "").strip()
    mode = data.get("mode", "url")
    max_notes = data.get("max_notes", 30)
    blogger_name = (data.get("blogger_name") or "").strip()
    deep_mode = data.get("deep_mode", "A")

    if mode != "self" and not url and not user_id:
        return jsonify({"ok": False, "error": "请提供博主 URL 或用户 ID"})

    params = {
        "url": url,
        "user_id": user_id,
        "mode": mode,
        "max_notes": max_notes,
        "blogger_name": blogger_name,
        "deep_mode": deep_mode,
    }

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {
        "id": task_id,
        "type": "pipeline",
        "queue": log_queue,
        "done": False,
        "exit_code": None,
        "result_files": {},
    }
    with _task_lock:
        _tasks[task_id] = task
    t = threading.Thread(target=_run_pipeline, args=(params, log_queue, task_id), daemon=True)
    t.start()
    return jsonify({"ok": True, "task_id": task_id})


@app.route("/api/start-crawl", methods=["POST", "GET"])
def start_crawl():
    if request.method == "GET":
        data = request.args
    else:
        data = request.get_json() or request.form
    url = (data.get("url") or "").strip()
    user_id = (data.get("user_id") or "").strip()
    mode = data.get("mode", "url")
    max_notes = data.get("max_notes", 30)

    python = sys.executable
    script = str(TOOLS_DIR / "博主蒸馏" / "scripts" / "spider_xhs_adapter.py")
    cmd = [python, script, "--output", str(DATA_DIR), "--max-notes", str(max_notes)]

    if url:
        cmd.extend(["--url", url])
    elif mode == "self":
        cmd.append("--self")
    elif mode == "id" and user_id:
        cmd.extend(["--user-id", user_id])
    else:
        return jsonify({"ok": False, "error": "请提供博主 URL 或用户 ID"})

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {"id": task_id, "type": "crawl", "cmd": cmd, "queue": log_queue, "done": False, "exit_code": None}
    with _task_lock:
        _tasks[task_id] = task
    t = threading.Thread(target=_run_subprocess, args=(cmd, log_queue, task_id), daemon=True)
    t.start()
    return jsonify({"ok": True, "task_id": task_id})


@app.route("/api/run-analysis", methods=["POST", "GET"])
def run_analysis():
    data = request.args if request.method == "GET" else (request.get_json() or request.form)
    details_file = (data.get("details_file") or "").strip()
    if not details_file or not os.path.isfile(details_file):
        found = _find_result_json()
        if not found:
            return jsonify({"ok": False, "error": "未找到采集数据文件，请先采集"})
        details_file = found

    python = sys.executable
    cmd = [python, str(TOOLS_DIR / "博主蒸馏" / "scripts" / "analyze.py"), details_file, "-o", str(DATA_DIR)]

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {"id": task_id, "type": "analysis", "cmd": cmd, "queue": log_queue, "done": False, "exit_code": None}
    with _task_lock:
        _tasks[task_id] = task
    t = threading.Thread(target=_run_subprocess, args=(cmd, log_queue, task_id), daemon=True)
    t.start()
    return jsonify({"ok": True, "task_id": task_id})


@app.route("/api/run-deep", methods=["POST", "GET"])
def run_deep():
    data = request.args if request.method == "GET" else (request.get_json() or request.form)
    blogger_name = (data.get("blogger_name") or "").strip()
    mode = data.get("mode", "A")
    if not blogger_name:
        return jsonify({"ok": False, "error": "请输入博主名称"})

    analysis_files = sorted(DATA_DIR.glob("*_analysis.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    if not analysis_files:
        return jsonify({"ok": False, "error": "未找到分析文件，请先运行分析"})

    analysis_file = str(analysis_files[0])
    details_files = sorted(DATA_DIR.glob("*_notes_details.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    details_file = str(details_files[0]) if details_files else ""

    python = sys.executable
    deep_out_dir = str(MATERIAL_SELF_DIR) if mode == "B" else str(MATERIAL_BLOGGER_DIR)
    os.makedirs(deep_out_dir, exist_ok=True)
    cmd = [
        python, str(TOOLS_DIR / "博主蒸馏" / "scripts" / "deep_analyze.py"),
        analysis_file, blogger_name,
        "-o", deep_out_dir,
        "--details", details_file,
        "--mode", mode,
    ]

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {"id": task_id, "type": "deep", "cmd": cmd, "queue": log_queue, "done": False, "exit_code": None}
    with _task_lock:
        _tasks[task_id] = task
    t = threading.Thread(target=_run_subprocess, args=(cmd, log_queue, task_id), daemon=True)
    t.start()
    return jsonify({"ok": True, "task_id": task_id})


# =============================================================
#  帖子深挖 — 路由
# =============================================================

def _scan_post_dive_history():
    """扫描帖子深挖历史记录"""
    pd_dir = OUTPUT_DIR / "帖子深挖"
    if not pd_dir.is_dir():
        return []
    import re
    records = []
    for md_file in sorted(pd_dir.glob("帖子分析_*.md"), key=lambda f: f.stat().st_mtime, reverse=True):
        note_id = md_file.stem.replace("帖子分析_", "")
        mtime = datetime.fromtimestamp(md_file.stat().st_mtime)
        title = ""
        try:
            content = md_file.read_text(encoding="utf-8")
            tm = re.search(r"\*\*标题\*\*:\s*(.+)", content)
            if tm:
                title = tm.group(1).strip()
            if not title or title == "无标题":
                am = re.search(r"\*\*作者\*\*:\s*(.+)", content)
                author = am.group(1).strip() if am else ""
                title = f"{author}的帖子" if author and author != "未知作者" else f"帖子 {note_id[:8]}"
        except Exception:
            title = f"帖子 {note_id[:8]}"
        url = f"https://www.xiaohongshu.com/explore/{note_id}"
        records.append({
            "note_id": note_id,
            "title": title,
            "url": url,
            "path": str(md_file),
            "time": mtime.strftime("%Y-%m-%d %H:%M"),
            "mtime_ts": md_file.stat().st_mtime,
        })
    return records


@app.route("/api/post-dive/history")
def post_dive_history():
    return jsonify({"ok": True, "records": _scan_post_dive_history()})


@app.route("/api/post-dive/read")
def post_dive_read():
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    pd_dir = (OUTPUT_DIR / "帖子深挖").resolve()
    if not str(full_path).startswith(str(pd_dir)):
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})
    try:
        content = full_path.read_text(encoding="utf-8")
    except Exception:
        return jsonify({"ok": False, "error": "read failed"})
    return jsonify({"ok": True, "content": content, "name": full_path.name})


@app.route("/api/start-post-dive", methods=["POST", "GET"])
def start_post_dive():
    data = request.args if request.method == "GET" else (request.get_json() or request.form)
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"ok": False, "error": "请提供帖子链接"})

    python = sys.executable
    run_script = str(TOOLS_DIR / "帖子深挖" / "run.py")
    # 同时支持 /item/ 和 /explore/ 两种 URL 格式
    note_id = "unknown"
    for pat in ["/item/", "/explore/"]:
        if pat in url:
            note_id = url.split(pat)[-1].split("?")[0]
            break
    output_dir = str(OUTPUT_DIR / "帖子深挖")
    os.makedirs(output_dir, exist_ok=True)
    cmd = [python, run_script, url, "-o", output_dir]

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {"id": task_id, "type": "post_dive", "cmd": cmd, "queue": log_queue, "done": False, "exit_code": None}
    with _task_lock:
        _tasks[task_id] = task
    t = threading.Thread(target=_run_subprocess, args=(cmd, log_queue, task_id), daemon=True)
    t.start()
    return jsonify({"ok": True, "task_id": task_id})


#  素材库 — 博主风格浏览
# =============================================================

_bloggers_cache = {"data": None, "mtime": 0}
_BLOGGERS_CACHE_TTL = 10


def _scan_bloggers(force=False):
    """扫描所有已蒸馏的博主，返回博主列表（带 10 秒缓存）"""
    import time as _time
    now = _time.time()
    if not force and _bloggers_cache["data"] is not None and (now - _bloggers_cache["mtime"]) < _BLOGGERS_CACHE_TTL:
        return _bloggers_cache["data"]
    bloggers = []
    seen = set()

    analysis_files = sorted(DATA_DIR.glob("*_analysis.json"), key=lambda f: f.stat().st_mtime, reverse=True)
    for af in analysis_files:
        name = af.name.replace("_analysis.json", "")
        if name in seen:
            continue
        seen.add(name)
        info = {"name": name, "files": {}, "notes_count": 0, "avg_likes": 0, "total_likes": 0}
        try:
            with open(af, "r", encoding="utf-8") as f:
                data = json.load(f)
            stats = data.get("stats", {})
            blogger = data.get("blogger", {})
            info["notes_count"] = stats.get("total", 0)
            info["avg_likes"] = stats.get("avg_likes", 0)
            info["total_likes"] = stats.get("total_likes", 0)
            info["total_collects"] = stats.get("total_collects", 0)
            info["total_comments"] = stats.get("total_comments", 0)
            info["category_stats"] = data.get("category_stats", {})
            info["blogger_desc"] = blogger.get("desc", "")
            info["blogger_id"] = blogger.get("bloggerId", "")
            info["tag_freq"] = data.get("tag_freq", {})
            info["notes"] = []
            for n in data.get("notes", []):
                info["notes"].append({
                    "id": n.get("id", ""),
                    "title": n.get("title", ""),
                    "likes": n.get("likes", 0),
                    "collects": n.get("collects", 0),
                    "comments_count": n.get("comments_count", 0),
                    "shares": n.get("shares", 0),
                    "tags": n.get("tags", []),
                    "category": n.get("category", ""),
                    "type": n.get("type", "normal"),
                    "time": n.get("time", ""),
                })
        except Exception:
            pass
        info["files"]["analysis_json"] = str(af)

        details_f = DATA_DIR / f"{name}_notes_details.json"
        if details_f.exists():
            info["files"]["details_json"] = str(details_f)

        for suffix, key in [
            ("_博主深度拆解.md", "deep_breakdown"),
            ("_内容公式总结.md", "formula_summary"),
            ("_选题素材库.md", "topic_library"),
            ("_数据底稿.md", "data_draft"),
            ("_AI蒸馏任务.md", "ai_task"),
            ("_全量笔记结构化分析.md", "structured_analysis"),
        ]:
            for root in [OUTPUT_DIR, MATERIAL_BLOGGER_DIR]:
                if root.is_dir():
                    matches = list(root.rglob(f"{name}{suffix}"))
                    if matches:
                        info["files"][key] = str(matches[0])
                        break

        # 从新结构（嵌套）/ 旧结构（扁平）两种方式找笔记目录
        notes_dir = None
        for root in [OUTPUT_DIR, MATERIAL_BLOGGER_DIR]:
            if root.is_dir():
                matches = sorted(root.rglob(f"{name}_笔记"), key=lambda p: len(str(p)))
                if matches:
                    notes_dir = matches[0]
                    break
        if notes_dir is not None and notes_dir.is_dir():
            info["files"]["notes_dir"] = str(notes_dir)
            info["posts"] = []
            for sub in sorted(notes_dir.iterdir()):
                if sub.is_dir():
                    body_file = sub / "正文.txt"
                    img_file = sub / "图片列表.txt"
                    post = {
                        "folder": sub.name,
                        "title": sub.name.split("_", 1)[-1] if "_" in sub.name else sub.name,
                    }
                    if body_file.exists():
                        try:
                            post["content"] = body_file.read_text(encoding="utf-8").strip()
                        except Exception:
                            pass
                    if img_file.exists():
                        try:
                            imgs = img_file.read_text(encoding="utf-8").strip().split("\n")
                            post["images"] = [l.strip() for l in imgs if l.strip()]
                        except Exception:
                            pass
                    info["posts"].append(post)

        bloggers.append(info)

    # 第二遍：扫描 素材库/博主风格/ 中的目录，补充没有 analysis.json 的博主
    if MATERIAL_BLOGGER_DIR.is_dir():
        for entry in sorted(MATERIAL_BLOGGER_DIR.iterdir()):
            if not entry.is_dir() or entry.name.startswith('_'):
                continue
            name = entry.name
            if name in seen:
                continue
            seen.add(name)

            info = {"name": name, "files": {}, "notes_count": 0, "avg_likes": 0, "total_likes": 0}

            # 尝试找笔记目录（兼容 笔记/ 和 {name}_笔记/ 两种命名）
            notes_dir = None
            for sub_name in [f"{name}_笔记", "笔记"]:
                p = entry / sub_name
                if p.is_dir():
                    notes_dir = p
                    break
            if notes_dir is not None:
                # 如果笔记目录的子目录里没有 正文.txt，可能是多了一层（如 笔记/薯岛_笔记/）
                inner = sorted(notes_dir.iterdir())
                if all(sub.is_dir() and not (sub / "正文.txt").exists() for sub in inner):
                    # 取第一个有子目录的级
                    for sub in inner:
                        deeper = list(sub.iterdir())
                        if any(d.is_dir() for d in deeper):
                            notes_dir = sub
                            break
                info["files"]["notes_dir"] = str(notes_dir)
                info["posts"] = []
                for sub in sorted(notes_dir.iterdir()):
                    if sub.is_dir():
                        body_file = sub / "正文.txt"
                        img_file = sub / "图片列表.txt"
                        post = {"folder": sub.name, "title": sub.name.split("_", 1)[-1] if "_" in sub.name else sub.name}
                        if body_file.exists():
                            try:
                                post["content"] = body_file.read_text(encoding="utf-8").strip()
                            except Exception:
                                pass
                        if img_file.exists():
                            try:
                                imgs = img_file.read_text(encoding="utf-8").strip().split("\n")
                                post["images"] = [l.strip() for l in imgs if l.strip()]
                            except Exception:
                                pass
                        info["posts"].append(post)
                info["notes_count"] = len(info.get("posts", []))

            # 读取 analysis_data.json 获取统计数据
            ad_file = entry / "analysis_data.json"
            if ad_file.exists():
                try:
                    with open(ad_file, "r", encoding="utf-8") as f:
                        ad = json.load(f)
                    if not info.get("notes_count"):
                        info["notes_count"] = ad.get("total_notes", 0)
                    if not info.get("total_likes"):
                        info["total_likes"] = ad.get("total_likes", 0)
                    if not info.get("avg_likes"):
                        info["avg_likes"] = round(ad.get("avg_likes", 0), 1)
                    info["files"]["analysis_data"] = str(ad_file)
                except Exception:
                    pass

            # 标记已存在的风格文件
            for skill_name in [f"{name}_创作指南.skill", "创作指南.skill"]:
                skill_dir = entry / skill_name
                if skill_dir.is_dir():
                    skill_md = skill_dir / "SKILL.md"
                    info["files"]["deep_breakdown"] = str(skill_md) if skill_md.exists() else str(skill_dir)
                    break

            # 扫描博主根目录下的 .md 文件（如 蒸馏报告.md、粗门问题.md 等）
            for md_file in sorted(entry.glob("*.md")):
                fn = md_file.name
                if "博主深度拆解" in fn or "深度拆解" in fn:
                    info["files"]["deep_breakdown"] = info["files"].get("deep_breakdown") or str(md_file)
                elif "内容公式" in fn or "公式总结" in fn:
                    info["files"]["formula_summary"] = str(md_file)
                elif "选题素材" in fn or "选题库" in fn:
                    info["files"]["topic_library"] = str(md_file)
                elif "数据底稿" in fn:
                    info["files"]["data_draft"] = str(md_file)
                elif "蒸馏报告" in fn:
                    info["files"]["distill_report"] = str(md_file)
                elif "帖子分析" in fn:
                    info["files"]["post_analysis"] = info["files"].get("post_analysis") or str(md_file)
                elif "诊断报告" in fn or "账号诊断" in fn:
                    info["files"]["diagnosis"] = str(md_file)
                else:
                    label = fn.replace(".md", "").replace(name, "").strip("_- ")
                    if label:
                        info["files"]["report_" + label] = str(md_file)
                    else:
                        info["files"]["report_" + fn.replace(".md", "")] = str(md_file)

            # 扫描 分析报告/ 子目录
            report_dir = entry / "分析报告"
            if report_dir.is_dir():
                for md_file in sorted(report_dir.rglob("*.md")):
                    fn = md_file.name
                    if "帖子分析" in fn:
                        info["files"]["post_analysis"] = info["files"].get("post_analysis") or str(md_file)
                    elif "博主深度拆解" in fn or "深度拆解" in fn:
                        info["files"]["deep_breakdown"] = info["files"].get("deep_breakdown") or str(md_file)
                    elif "蒸馏报告" in fn:
                        info["files"]["distill_report"] = str(md_file)
                    elif "诊断报告" in fn or "账号诊断" in fn:
                        info["files"]["diagnosis"] = str(md_file)
                    else:
                        info["files"]["report_" + fn.replace(".md", "")] = str(md_file)
                for html_file in sorted(report_dir.rglob("*.html")):
                    fn = html_file.name
                    if "蒸馏报告" in fn:
                        info["files"]["distill_report_html"] = str(html_file)

            # 扫描 _过程文件/原始素材/ 中的关联文件
            process_dir = MATERIAL_BLOGGER_DIR / "_过程文件" / "原始素材"
            if process_dir.is_dir():
                for md_file in sorted(process_dir.glob(f"{name}_*.md")):
                    fn = md_file.name
                    if "博主深度拆解" in fn:
                        info["files"]["deep_breakdown"] = info["files"].get("deep_breakdown") or str(md_file)
                    elif "内容公式" in fn or "公式总结" in fn:
                        info["files"]["formula_summary"] = info["files"].get("formula_summary") or str(md_file)
                    elif "选题素材" in fn or "选题库" in fn:
                        info["files"]["topic_library"] = info["files"].get("topic_library") or str(md_file)
                    elif "数据底稿" in fn:
                        info["files"]["data_draft"] = info["files"].get("data_draft") or str(md_file)
                    elif "全量笔记结构化分析" in fn:
                        info["files"]["structured_analysis"] = info["files"].get("structured_analysis") or str(md_file)

            bloggers.append(info)

    _bloggers_cache["data"] = bloggers
    _bloggers_cache["mtime"] = _time.time()
    return bloggers


@app.route("/api/bloggers")
def api_bloggers():
    return jsonify({"ok": True, "bloggers": _scan_bloggers()})


@app.route("/api/blogger/<name>")
def api_blogger_detail(name):
    bloggers = _scan_bloggers()
    for b in bloggers:
        if b["name"] == name:
            return jsonify({"ok": True, "blogger": b})
    return jsonify({"ok": False, "error": "博主不存在"}), 404


@app.route("/api/blogger-file")
def api_blogger_file():
    """读取博主的某个产出文件"""
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})
    try:
        content = full_path.read_text(encoding="utf-8")
    except Exception:
        return jsonify({"ok": False, "error": "read failed"})
    return jsonify({"ok": True, "content": content, "name": full_path.name})


@app.route("/api/blogger-html")
def api_blogger_html():
    filepath = request.args.get("path", "")
    if not filepath:
        return "no path", 400
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return "access denied", 403
    if not full_path.is_file():
        return "file not found", 404
    try:
        content = full_path.read_text(encoding="utf-8")
    except Exception:
        return "read failed", 500
    return content


def _collect_blogger_files(blogger_name):
    """收集博主的所有产出文件，返回 {label: filepath} 字典"""
    files = {}
    name = blogger_name.strip()
    if not name:
        return files

    file_patterns = [
        ("博主深度拆解", "_博主深度拆解.md"),
        ("内容公式总结", "_内容公式总结.md"),
        ("选题素材库", "_选题素材库.md"),
        ("数据底稿", "_数据底稿.md"),
        ("AI蒸馏任务", "_AI蒸馏任务.md"),
        ("全量笔记结构化分析", "_全量笔记结构化分析.md"),
        ("蒸馏报告", "_蒸馏报告.md"),
        ("诊断报告", "_诊断报告.md"),
    ]

    search_roots = [OUTPUT_DIR, MATERIAL_BLOGGER_DIR]
    for label, suffix in file_patterns:
        for root in search_roots:
            if root.is_dir():
                matches = list(root.rglob(f"{name}{suffix}"))
                if matches:
                    files[label] = str(matches[0])
                    break

    for root in search_roots:
        if root.is_dir():
            for entry in sorted(root.iterdir()):
                if not entry.is_dir():
                    continue
                if entry.name.startswith(name):
                    skill_dir = entry / "创作指南.skill"
                    if skill_dir.is_dir():
                        skill_md = skill_dir / "SKILL.md"
                        if skill_md.exists():
                            files["创作指南"] = str(skill_md)
                    report_dir = entry / "分析报告"
                    if report_dir.is_dir():
                        for md_file in sorted(report_dir.rglob("*.md")):
                            fn = md_file.name
                            if "诊断" in fn and "诊断报告" not in files:
                                files["诊断报告"] = str(md_file)
                            elif "蒸馏报告" in fn and "蒸馏报告" not in files:
                                files["蒸馏报告"] = str(md_file)
                            elif "深度拆解" in fn and "博主深度拆解" not in files:
                                files["博主深度拆解"] = str(md_file)
                            elif any(k not in files for k in ["博主深度拆解", "内容公式总结", "选题素材库"]):
                                label = fn.replace(".md", "").replace(name, "").strip("_- ")
                                if label:
                                    files[label] = str(md_file)

    return files


def _build_export_html(blogger_name, files):
    """将博主的所有产出文件合并为一个 HTML 文档"""
    import html as html_mod
    sections = []
    for label, filepath in files.items():
        try:
            content = Path(filepath).read_text(encoding="utf-8")
        except Exception:
            content = f"（读取失败：{filepath}）"
        sections.append(f'<h2 style="color:#ff6b6b;border-bottom:2px solid #ff6b6b;padding-bottom:8px;margin-top:40px">{html_mod.escape(label)}</h2>')
        sections.append(f'<div style="white-space:pre-wrap;line-height:1.8;font-size:14px">{html_mod.escape(content)}</div>')

    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{html_mod.escape(blogger_name)} - 博主分析报告</title>
<style>
body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif; max-width: 900px; margin: 0 auto; padding: 40px 20px; color: #333; }}
h1 {{ color: #ff6b6b; border-bottom: 3px solid #ff6b6b; padding-bottom: 12px; }}
h2 {{ color: #ff6b6b; border-bottom: 2px solid #ff6b6b; padding-bottom: 8px; margin-top: 40px; }}
.meta {{ color: #999; font-size: 13px; margin-bottom: 30px; }}
@media print {{ body {{ padding: 20px; }} h2 {{ page-break-before: auto; }} }}
</style>
</head>
<body>
<h1>{html_mod.escape(blogger_name)} - 博主分析报告</h1>
<div class="meta">导出时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M')} | 小红书工具箱</div>
{''.join(sections)}
</body>
</html>"""
    return html_content


def _build_export_docx(blogger_name, files):
    """将博主的所有产出文件合并为一个 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    title = doc.add_heading(blogger_name + ' - 博主分析报告', level=0)
    title.runs[0].font.color.rgb = RGBColor(0xFF, 0x6B, 0x6B)

    meta = doc.add_paragraph(f'导出时间：{__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")} | 小红书工具箱')
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for label, filepath in files.items():
        try:
            content = Path(filepath).read_text(encoding="utf-8")
        except Exception:
            content = f"（读取失败：{filepath}）"

        h = doc.add_heading(label, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x6B)

        for line in content.split("\n"):
            line = line.strip()
            if line.startswith("# "):
                h2 = doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=4)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                doc.add_paragraph(line)

    return doc


@app.route("/api/export/blogger", methods=["POST"])
def export_blogger():
    """导出博主分析报告为 HTML / Word / PDF"""
    data = request.get_json() or request.form
    blogger_name = (data.get("name") or "").strip()
    fmt = (data.get("format") or "html").strip().lower()

    if not blogger_name:
        return jsonify({"ok": False, "error": "请提供博主名称"})

    files = _collect_blogger_files(blogger_name)
    if not files:
        return jsonify({"ok": False, "error": f"未找到博主「{blogger_name}」的分析结果"})

    if fmt == "html":
        html_content = _build_export_html(blogger_name, files)
        export_dir = OUTPUT_DIR / "导出报告"
        export_dir.mkdir(parents=True, exist_ok=True)
        export_path = export_dir / f"{blogger_name}_分析报告.html"
        export_path.write_text(html_content, encoding="utf-8")
        return jsonify({"ok": True, "path": str(export_path), "format": "html"})

    elif fmt == "word":
        try:
            from docx import Document
        except ImportError:
            return jsonify({"ok": False, "error": "缺少 python-docx 库，请运行: python -m pip install python-docx"})
        doc = _build_export_docx(blogger_name, files)
        export_dir = OUTPUT_DIR / "导出报告"
        export_dir.mkdir(parents=True, exist_ok=True)
        export_path = export_dir / f"{blogger_name}_分析报告.docx"
        doc.save(str(export_path))
        return jsonify({"ok": True, "path": str(export_path), "format": "word"})

    elif fmt == "pdf":
        html_content = _build_export_html(blogger_name, files)
        export_dir = OUTPUT_DIR / "导出报告"
        export_dir.mkdir(parents=True, exist_ok=True)
        html_path = export_dir / f"{blogger_name}_分析报告_打印.html"
        html_path.write_text(html_content, encoding="utf-8")
        return jsonify({"ok": True, "path": str(html_path), "format": "pdf", "hint": "已生成 HTML 报告，请在浏览器中打开后使用 Ctrl+P 打印为 PDF"})

    else:
        return jsonify({"ok": False, "error": f"不支持的格式: {fmt}"})


@app.route("/api/export/download")
def export_download():
    """下载导出的文件"""
    filepath = request.args.get("path", "")
    if not filepath:
        return "no path", 400
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return "access denied", 403
    if not full_path.is_file():
        return "file not found", 404
    from flask import send_file
    return send_file(str(full_path), as_attachment=True)


# =============================================================
#  素材库 — 路由
# =============================================================

_MATERIAL_ROOTS = {
    "博主风格": str(MATERIAL_BLOGGER_DIR.resolve()),
    "自我参数": str(MATERIAL_SELF_DIR.resolve()),
}


@app.route("/api/material-lib/tree")
def material_lib_tree():
    """返回素材库两个分类的目录树"""
    sections = {}
    for section_name, root_path in _MATERIAL_ROOTS.items():
        root = Path(root_path)
        if root.is_dir():
            sections[section_name] = {
                "root": str(root),
                "tree": _tree_walk(root),
            }
        else:
            sections[section_name] = {"root": str(root), "tree": []}
    return jsonify({"ok": True, "sections": sections})


@app.route("/api/material-lib/read")
def material_lib_read():
    """读取素材库中的文件内容（安全检查：必须在任一素材库目录内）"""
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()

    allowed = False
    for root_path in _MATERIAL_ROOTS.values():
        if str(full_path).startswith(root_path):
            allowed = True
            break
    if not allowed:
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})

    try:
        content = full_path.read_text(encoding="utf-8")
    except Exception:
        return jsonify({"ok": False, "error": "read failed (binary?)"})

    ext = full_path.suffix.lower()
    return jsonify({
        "ok": True,
        "content": content,
        "name": full_path.name,
        "ext": ext,
        "section": next((s for s, p in _MATERIAL_ROOTS.items() if str(full_path).startswith(p)), ""),
    })


@app.route("/api/material-lib/write", methods=["POST"])
def material_lib_write():
    """写入素材库文件（安全检查：必须在自我参数目录内）"""
    body = request.get_json(force=True)
    filepath = body.get("path", "")
    content = body.get("content", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    self_dir = MATERIAL_SELF_DIR.resolve()
    if not str(full_path).startswith(str(self_dir)):
        return jsonify({"ok": False, "error": "only self-params files can be edited"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})
    try:
        full_path.write_text(content, encoding="utf-8")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})
    return jsonify({"ok": True})


# =============================================================
#  共享路由
# =============================================================

@app.route("/")
def index():
    return render_template("index.html", output_dir=str(OUTPUT_DIR.resolve()))


@app.route("/api/task/<task_id>/stream")
def task_stream(task_id: str):
    with _task_lock:
        task = _tasks.get(task_id)
    if not task:
        return jsonify({"error": "task not found"}), 404

    log_queue = task["queue"]

    def generate():
        while True:
            line = log_queue.get()
            if line is None:
                with _task_lock:
                    t = _tasks.get(task_id, {})
                rf = t.get("result_files", {})
                if not rf and t.get("done") and t.get("exit_code") == 0:
                    tt = t.get("type", "")
                    if tt == "keyword_search":
                        kw = t.get("keyword", "")
                        files = sorted(SEARCH_OUTPUT_DIR.glob(f"{kw}_搜索结果_*.json"),
                                       key=lambda f: f.stat().st_mtime, reverse=True)
                        if files:
                            rf["search_result"] = str(files[0])
                            with _task_lock:
                                _tasks[task_id]["result_files"] = rf
                payload = json.dumps({
                    "exit_code": t.get("exit_code", -1),
                    "result_files": rf,
                    "type": t.get("type", ""),
                })
                yield f"event: done\ndata: {payload}\n\n"
                break
            stripped = line.rstrip("\n")
            if stripped.startswith("__PHASE__:"):
                phase_num = stripped.split(":")[1]
                yield f"event: phase\ndata: {phase_num}\n\n"
            elif stripped.startswith("__PHASE_DONE__:"):
                phase_num = stripped.split(":")[1]
                yield f"event: phase_done\ndata: {phase_num}\n\n"
            elif stripped.startswith("[ERROR]"):
                yield f"event: error_event\ndata: {json.dumps(stripped)}\n\n"
                yield f"data: {json.dumps(stripped)}\n\n"
            else:
                encoded = json.dumps(stripped)
                yield f"data: {encoded}\n\n"

    return Response(generate(), mimetype="text/event-stream")


@app.route("/api/task/<task_id>/status")
def task_status(task_id: str):
    with _task_lock:
        task = _tasks.get(task_id)
    if not task:
        return jsonify({"error": "not found"}), 404

    done = task.get("done", False)
    exit_code = task.get("exit_code")
    # 优先用任务存储的结果文件，再按类型回退查找
    result_files = task.get("result_files", {})
    if done and not result_files:
        t = task["type"]
        if t == "crawl":
            f = _find_result_json()
            if f:
                result_files["details_json"] = f
        elif t == "analysis":
            af = _find_analysis_json()
            if af:
                result_files["analysis_json"] = af
        elif t in ("deep",):
            result_files = _find_deep_output()
        elif t == "post_dive":
            pd_dir = OUTPUT_DIR / "帖子深挖"
            md_files = sorted(pd_dir.glob("*.md"), key=lambda f: f.stat().st_mtime, reverse=True)
            if md_files:
                result_files["report_md"] = str(md_files[0])
        elif t == "keyword_search":
            kw = task.get("keyword", "")
            files = sorted(SEARCH_OUTPUT_DIR.glob(f"{kw}_搜索结果_*.json"),
                           key=lambda f: f.stat().st_mtime, reverse=True)
            if files:
                result_files["search_result"] = str(files[0])

    return jsonify({
        "ok": True, "task_id": task_id, "done": done,
        "exit_code": exit_code, "type": task["type"], "result_files": result_files,
    })


@app.route("/api/task/<task_id>/stop", methods=["POST"])
def task_stop(task_id: str):
    """终止正在运行的任务"""
    # 设置停止标志，防止管道在阶段间隙继续
    _get_stop_event(task_id).set()

    with _procs_lock:
        proc = _running_procs.get(task_id)

    if proc:
        try:
            proc.kill()
            # 关闭管道强制 readline 立即返回
            if proc.stdout:
                proc.stdout.close()
            proc.wait(timeout=5)
        except Exception:
            pass

        with _procs_lock:
            _running_procs.pop(task_id, None)

    # 即使没有找到进程（阶段间隙），也标记任务为已完成
    with _task_lock:
        if task_id in _tasks:
            _tasks[task_id]["done"] = True
            _tasks[task_id]["exit_code"] = -1

    return jsonify({"ok": True, "message": "任务已终止"})


@app.route("/api/download")
def download_file():
    filepath = request.args.get("path", "")
    if not filepath:
        return "no path", 400
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return "access denied", 403
    if not full_path.is_file():
        return "File not found", 404
    filename = full_path.name
    return send_file(str(full_path), as_attachment=True, download_name=filename)


@app.route("/api/find-results")
def find_results():
    result = {"details_files": [], "analysis_files": [], "data_drafts": [], "ai_tasks": []}
    for f in sorted(DATA_DIR.glob("*_notes_details.json"), key=lambda f: f.stat().st_mtime, reverse=True):
        result["details_files"].append(str(f))
    for f in sorted(DATA_DIR.glob("*_analysis.json"), key=lambda f: f.stat().st_mtime, reverse=True):
        result["analysis_files"].append(str(f))
    for f in sorted(OUTPUT_DIR.glob("*_数据底稿.md"), key=lambda f: f.stat().st_mtime, reverse=True):
        result["data_drafts"].append(str(f))
    for f in sorted(OUTPUT_DIR.glob("*_AI蒸馏任务.md"), key=lambda f: f.stat().st_mtime, reverse=True):
        result["ai_tasks"].append(str(f))
    return jsonify(result)


@app.route("/api/self-account")
def self_account():
    """返回当前 Cookie 对应的登录账号信息"""
    try:
        scripts_dir = TOOLS_DIR / "博主蒸馏" / "scripts"
        helper = scripts_dir / "get_self_info.py"
        python = sys.executable
        r = subprocess.run(
            [python, str(helper)],
            capture_output=True, text=True, timeout=30,
            env={**os.environ, "PYTHONUTF8": "1", "PYTHONIOENCODING": "utf-8"},
        )
        out = r.stdout.strip()
        if out:
            return jsonify(json.loads(out))
        return jsonify({"ok": False, "error": f"no output (stderr: {r.stderr[:200]})"})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/cookie-info")
def cookie_info():
    """返回当前 Cookie 状态"""
    env_path = SPIDER_DIR / ".env"
    if not env_path.is_file():
        return jsonify({"ok": True, "exists": False, "length": 0, "preview": ""})
    try:
        content = env_path.read_text(encoding="utf-8")
        import re
        m = re.search(r'COOKIES\s*=\s*(.*)', content)
        if m:
            val = m.group(1).strip().strip("'\"").strip()
            preview = val[:30] + "..." if len(val) > 30 else val
            return jsonify({"ok": True, "exists": True, "length": len(val), "preview": preview})
        return jsonify({"ok": True, "exists": False, "length": 0, "preview": ""})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/update-cookie", methods=["POST"])
def update_cookie():
    """更新 .env 文件中的 Cookie"""
    data = request.get_json() or request.form
    new_cookie = (data.get("cookie") or "").strip().strip("'").strip('"')
    if not new_cookie:
        return jsonify({"ok": False, "error": "Cookie 不能为空"})
    if len(new_cookie) < 20:
        return jsonify({"ok": False, "error": "Cookie 太短，请复制完整的 Cookie 字符串"})

    env_path = SPIDER_DIR / ".env"
    try:
        if env_path.is_file():
            content = env_path.read_text(encoding="utf-8")
            import re
            if re.search(r'^COOKIES\s*=', content, re.MULTILINE):
                content = re.sub(
                    r'^COOKIES\s*=.*$',
                    f'COOKIES={new_cookie}',
                    content,
                    count=1,
                    flags=re.MULTILINE,
                )
            else:
                content += f'\nCOOKIES={new_cookie}\n'
        else:
            content = f'# 小红书 Cookie\nCOOKIES={new_cookie}\n'

        env_path.write_text(content, encoding="utf-8")
        return jsonify({"ok": True, "message": "Cookie 已更新", "preview": new_cookie[:30] + "..."})
    except Exception as e:
        return jsonify({"ok": False, "error": f"写入失败: {e}"})


@app.route("/api/check-env", methods=["GET"])
def check_env():
    checks = {}
    checks["python"] = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
    checks["spider_xhs"] = str(SPIDER_DIR.resolve()) if SPIDER_DIR.is_dir() else "NOT_FOUND"
    checks["env_file"] = "FOUND" if (SPIDER_DIR / ".env").is_file() else "NOT_FOUND"
    try:
        r = subprocess.run(["node", "--version"], capture_output=True, text=True, timeout=5)
        checks["nodejs"] = r.stdout.strip() if r.returncode == 0 else "NOT_FOUND"
    except Exception:
        checks["nodejs"] = "NOT_FOUND"
    for pkg in ["requests", "flask", "execjs"]:
        try:
            __import__(pkg)
            checks[f"pkg_{pkg}"] = "OK"
        except ImportError:
            checks[f"pkg_{pkg}"] = "MISSING"
    return jsonify({"ok": True, "checks": checks})


@app.route("/api/health")
def health_check():
    return jsonify({"ok": True, "status": "running", "port": int(os.environ.get("XHS_PORT", 5001))})


# =============================================================
#  历史结果 & 文件预览
# =============================================================

def _file_info(f: Path) -> dict:
    return {
        "name": f.name,
        "path": str(f),
        "size": f.stat().st_size,
        "mtime": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
        "ext": f.suffix.lower(),
    }


@app.route("/api/history-results")
def history_results():
    """扫描所有产出目录，按工具分类返回结果列表"""
    results = {
        "博主分析": {"采集数据": [], "分析结果": [], "深度蒸馏": []},
        "帖子深挖": {"分析报告": []},
    }

    for f in sorted(DATA_DIR.glob("*_notes_details.json"), key=lambda f: f.stat().st_mtime, reverse=True):
        results["博主分析"]["采集数据"].append(_file_info(f))
    for f in sorted(DATA_DIR.glob("*_analysis.json"), key=lambda f: f.stat().st_mtime, reverse=True):
        results["博主分析"]["分析结果"].append(_file_info(f))
    for pattern in ["*_数据底稿.md", "*_AI蒸馏任务.md", "*_分析报告.html", "*_创作指南.skill"]:
        for root in [OUTPUT_DIR, OUTPUT_DIR / "博主分析", MATERIAL_BLOGGER_DIR]:
            if root.is_dir():
                for f in sorted(root.glob(pattern), key=lambda f: f.stat().st_mtime, reverse=True):
                    if not any(f.name == existing["name"] for existing in results["博主分析"]["深度蒸馏"]):
                        results["博主分析"]["深度蒸馏"].append(_file_info(f))

    for f in sorted((OUTPUT_DIR / "帖子深挖").glob("*.md"), key=lambda f: f.stat().st_mtime, reverse=True):
        info = _file_info(f)
        # 从文件内容中提取帖子标题
        try:
            text = f.read_text(encoding="utf-8", errors="replace")
            m = re.search(r'\*\*标题\*\*:\s*(.+)', text)
            if m:
                info["display_name"] = m.group(1).strip()
        except Exception:
            pass
        results["帖子深挖"]["分析报告"].append(info)

    return jsonify({"ok": True, "results": results})


@app.route("/api/delete-file", methods=["POST"])
def delete_file():
    """删除指定的产出文件"""
    data = request.get_json() or request.form
    filepath = (data.get("path") or "").strip()
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})
    try:
        os.remove(full_path)
        return jsonify({"ok": True, "message": f"已删除 {full_path.name}"})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/file-preview")
def file_preview():
    """读取文件内容用于预览"""
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    # 安全检查：必须在项目目录内
    if not str(full_path).startswith(str(ROOT.resolve())):
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})

    ext = full_path.suffix.lower()
    try:
        raw = full_path.read_bytes()
        content = raw.decode("utf-8")
    except UnicodeDecodeError:
        return jsonify({"ok": False, "error": "binary file", "binary": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})

    return jsonify({
        "ok": True,
        "content": content,
        "name": full_path.name,
        "ext": ext,
    })


# =============================================================
#  关键词搜索
# =============================================================

SEARCH_OUTPUT_DIR = OUTPUT_DIR / "关键词搜索"
SEARCH_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _scan_search_history():
    """扫描关键词搜索历史"""
    if not SEARCH_OUTPUT_DIR.is_dir():
        return []
    records = []
    for json_file in sorted(SEARCH_OUTPUT_DIR.glob("*_搜索结果_*.json"), key=lambda f: f.stat().st_mtime, reverse=True):
        keyword = ""
        total = 0
        search_time = ""
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            keyword = data.get("keyword", "")
            total = data.get("total_notes", 0)
            search_time = data.get("search_time", "")
        except Exception:
            keyword = json_file.stem.split("_搜索结果_")[0]
        mtime = datetime.fromtimestamp(json_file.stat().st_mtime)
        records.append({
            "keyword": keyword,
            "total": total,
            "search_time": search_time or mtime.strftime("%Y-%m-%d %H:%M"),
            "path": str(json_file),
            "mtime_ts": json_file.stat().st_mtime,
        })
    return records


@app.route("/api/keyword-search/start", methods=["POST"])
def keyword_search_start():
    """启动关键词搜索任务"""
    data = request.get_json() or request.form
    keyword = (data.get("keyword") or "").strip()
    if not keyword:
        return jsonify({"ok": False, "error": "请输入搜索关键词"})

    max_notes = min(int(data.get("max_notes") or 20), 50)
    time_range = data.get("time_range") or "all"
    sort_type = data.get("sort_type") or "general"
    fetch_comments = data.get("fetch_comments", True)

    python = sys.executable
    cmd = [
        python,
        str(TOOLS_DIR / "博主蒸馏" / "scripts" / "keyword_search.py"),
        "--keyword", keyword,
        "--max-notes", str(max_notes),
        "--time-range", time_range,
        "--sort", sort_type,
        "--output", str(SEARCH_OUTPUT_DIR),
    ]
    if not fetch_comments:
        cmd.append("--no-comments")

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {
        "id": task_id,
        "type": "keyword_search",
        "cmd": cmd,
        "queue": log_queue,
        "done": False,
        "exit_code": None,
        "keyword": keyword,
        "result_files": {},
    }
    with _task_lock:
        _tasks[task_id] = task

    def _on_search_done(task_id, exit_code):
        if exit_code == 0:
            files = sorted(SEARCH_OUTPUT_DIR.glob(f"{keyword}_搜索结果_*.json"),
                           key=lambda f: f.stat().st_mtime, reverse=True)
            if files:
                result_path = str(files[0])
                with _task_lock:
                    _tasks[task_id]["result_files"]["search_result"] = result_path
                # 同时写入 SQLite
                try:
                    from db import save_search_session
                    content = files[0].read_text(encoding="utf-8")
                    data = json.loads(content)
                    save_search_session(keyword, data, result_path)
                except Exception as e:
                    print(f"[DB] 保存搜索记录失败: {e}")

    t = threading.Thread(target=_run_subprocess, args=(cmd, log_queue, task_id), kwargs={"on_done": _on_search_done}, daemon=True)
    t.start()

    return jsonify({"ok": True, "task_id": task_id})


@app.route("/api/keyword-search/history")
def keyword_search_history():
    """获取搜索历史"""
    return jsonify({"ok": True, "records": _scan_search_history()})


@app.route("/api/keyword-search/read")
def keyword_search_read():
    """读取搜索结果"""
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})
    try:
        content = full_path.read_text(encoding="utf-8")
        data = json.loads(content)
        return jsonify({"ok": True, "data": data})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


# =============================================================
#  粗趣 RAG 问答
# =============================================================

from 工具类.rag_engine import get_engine


@app.route("/api/rag/status")
def rag_status():
    """RAG 引擎状态"""
    engine = get_engine()
    st = engine.status()
    return jsonify({"ok": True, **st})


@app.route("/api/rag/load", methods=["POST"])
def rag_load():
    """加载/重新加载知识库"""
    engine = get_engine()
    ok = engine.load_knowledge_base(force=True)
    if ok:
        return jsonify({"ok": True, "message": f"知识库加载成功，共 {len(engine.chunks)} 个片段"})
    return jsonify({"ok": False, "error": "知识库加载失败"})


@app.route("/api/rag/query", methods=["POST"])
def rag_query():
    """RAG 问答"""
    data = request.get_json() or request.form
    question = (data.get("question") or "").strip()
    if not question:
        return jsonify({"ok": False, "error": "请输入问题"})

    engine = get_engine()
    if not engine.ready:
        ok = engine.load_knowledge_base()
        if not ok:
            return jsonify({"ok": False, "error": "知识库加载失败"})

    result = engine.query(question)
    return jsonify(result)


# =============================================================
#  图片缓存代理
# =============================================================

IMAGE_CACHE_DIR = Path(__file__).resolve().parent / "output" / "image_cache"


@app.route("/api/image-proxy")
def image_proxy():
    """代理小红书图片，缓存到本地"""
    url = request.args.get("url", "")
    if not url:
        return "", 400

    # 先从缓存取
    from image_cache_helper import get_cached_path, cache_image
    cached = get_cached_path(url)
    if cached:
        return send_from_directory(os.path.dirname(cached), os.path.basename(cached))

    cached = cache_image(url)
    if cached:
        return send_from_directory(os.path.dirname(cached), os.path.basename(cached))

    # 最后尝试直接 302 跳转
    return "", 404


# =============================================================
#  DB 统计
# =============================================================

@app.route("/api/db-stats")
def db_stats():
    try:
        from db import query_search_stats, query_history
        stats = query_search_stats()
        recent = query_history(limit=10)
        return jsonify({"ok": True, "stats": stats, "recent": recent})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


# =============================================================
#  关键词搜索 — Excel 导出
# =============================================================

@app.route("/api/keyword-search/export")
def keyword_search_export():
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"}), 400
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return jsonify({"ok": False, "error": "access denied"}), 403
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"}), 404

    try:
        # 复用 spider_xhs 的 save_to_xlsx 逻辑
        sys.path.insert(0, str(SPIDER_DIR))
        from xhs_utils.data_util import save_to_xlsx
        with open(full_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        results = data.get("results", data.get("data", {}).get("results", []))
        xlsx_path = str(full_path) + ".xlsx"
        save_to_xlsx(results, xlsx_path)
        return send_file(xlsx_path, as_attachment=True, download_name=Path(xlsx_path).name)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# (热门推荐功能已移除)


# =============================================================
#  用户搜索
# =============================================================

USERSEARCH_OUTPUT_DIR = ROOT / "output" / "用户搜索"
USERSEARCH_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


@app.route("/api/user-search/start", methods=["POST"])
def user_search_start():
    data = request.get_json() or request.form
    keyword = (data.get("keyword") or "").strip()
    max_users = min(int(data.get("max_users") or 10), 30)
    if not keyword:
        return jsonify({"ok": False, "error": "请输入搜索关键词"})

    python = sys.executable
    script = str(TOOLS_DIR / "博主蒸馏" / "scripts" / "user_search.py")
    cmd = [python, script, "--keyword", keyword, "--max-users", str(max_users), "--output", str(USERSEARCH_OUTPUT_DIR)]

    task_id = _next_id()
    log_queue = queue.Queue()
    task = {
        "id": task_id, "type": "user_search", "cmd": cmd,
        "queue": log_queue, "done": False, "exit_code": None,
        "keyword": keyword,
        "result_files": {},
    }
    with _task_lock:
        _tasks[task_id] = task

    t = threading.Thread(target=_run_subprocess, args=(cmd, log_queue, task_id), daemon=True)
    t.start()
    return jsonify({"ok": True, "task_id": task_id})


@app.route("/api/user-search/read")
def user_search_read():
    filepath = request.args.get("path", "")
    if not filepath:
        return jsonify({"ok": False, "error": "no path"})
    full_path = Path(filepath).resolve()
    if not str(full_path).startswith(str(ROOT.resolve())):
        return jsonify({"ok": False, "error": "access denied"})
    if not full_path.is_file():
        return jsonify({"ok": False, "error": "file not found"})
    try:
        content = full_path.read_text(encoding="utf-8")
        return jsonify({"ok": True, "data": json.loads(content)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/user-search/history")
def user_search_history():
    records = []
    for f in sorted(USERSEARCH_OUTPUT_DIR.glob("usersearch_*.json"),
                     key=lambda f: f.stat().st_mtime, reverse=True):
        keyword = ""
        total = 0
        search_time = ""
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            keyword = data.get("keyword", "")
            total = data.get("total", 0)
            search_time = data.get("search_time", "")
        except Exception:
            keyword = f.stem.replace("usersearch_", "").rsplit("_", 1)[0]
        from datetime import datetime
        mtime = datetime.fromtimestamp(f.stat().st_mtime)
        records.append({
            "keyword": keyword,
            "total": total,
            "search_time": search_time or mtime.strftime("%Y-%m-%d %H:%M"),
            "path": str(f),
            "mtime_ts": f.stat().st_mtime,
        })
    return jsonify({"ok": True, "records": records})


# =============================================================
#  启动
# =============================================================

if __name__ == "__main__":
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")

    PORT = int(os.environ.get("XHS_PORT", 5001))

    # 自动释放端口：杀掉占用 5001 的旧进程
    if sys.platform == "win32":
        try:
            r = subprocess.run(
                ["netstat", "-ano"], capture_output=True, text=True, timeout=5
            )
            for line in r.stdout.splitlines():
                if f":{PORT}" in line and "LISTENING" in line:
                    parts = line.strip().split()
                    pid = parts[-1]
                    subprocess.run(
                        ["taskkill", "/F", "/PID", pid],
                        capture_output=True, timeout=5,
                    )
                    print(f"  🧹 已释放端口 {PORT} (PID {pid})")
                    break
        except Exception:
            pass

    print("=" * 50)
    print(" 小红书工具箱 — 统一 Web 入口")
    print("=" * 50)
    print(f" 启动: http://localhost:{PORT}")
    print(f" 工具: 博主分析 / 帖子深挖")
    print(f" 素材库: {MATERIAL_DIR.resolve()}")
    print(f" 产出: {OUTPUT_DIR.resolve()}")
    print("=" * 50)

    def open_browser():
        import urllib.request
        for i in range(30):
            try:
                urllib.request.urlopen(f"http://localhost:{PORT}/api/health", timeout=1)
                import webbrowser
                webbrowser.open(f"http://localhost:{PORT}")
                print(f"  ✅ 浏览器已自动打开")
                return
            except Exception:
                time.sleep(0.5)
        print(f"  ⚠️ 浏览器未自动打开，请手动访问 http://localhost:{PORT}")

    threading.Thread(target=open_browser, daemon=True).start()

    app.run(host="0.0.0.0", port=PORT, debug=False, threaded=True)
